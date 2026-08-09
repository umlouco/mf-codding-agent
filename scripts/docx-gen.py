#!/usr/bin/env python3
"""docx-gen.py — Assemble a .docx from a docgen manifest JSON.

Reads a manifest (produced by the Go docgen tool) containing a markdown file
path and an optional array of screenshot references, converts the markdown into
a well-formatted Word document with headings, code blocks, embedded images and
captions, and a table of contents.

Usage:
    python scripts/docx-gen.py --manifest path/to/manifest.json --output out.docx
    python scripts/docx-gen.py --help
"""

import argparse
import json
import os
import re
import sys
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import _TableStyle


# ---------------------------------------------------------------------------
# Manifest helpers
# ---------------------------------------------------------------------------

def load_manifest(path: str) -> dict[str, Any]:
    """Read and validate the manifest JSON."""
    if not os.path.isfile(path):
        sys.exit(f"manifest not found: {path}")

    with open(path, "r", encoding="utf-8") as fh:
        try:
            data = json.load(fh)
        except json.JSONDecodeError as exc:
            sys.exit(f"invalid manifest JSON in {path}: {exc}")

    if "markdown" not in data:
        sys.exit("manifest is missing required 'markdown' field")

    return data


def resolve_path(manifest_dir: str, rel: str) -> str:
    """Resolve a path that may be relative to the manifest directory."""
    if os.path.isabs(rel):
        return rel
    return os.path.normpath(os.path.join(manifest_dir, rel))


def normalise_screenshots(raw: Any, manifest_dir: str) -> list[dict[str, str]]:
    """Normalise the screenshots array into a list of {path, caption} dicts.

    Each entry may be a plain string (path only) or a dict with 'path' and an
    optional 'caption'.
    """
    if not raw or not isinstance(raw, list):
        return []

    shots: list[dict[str, str]] = []
    for entry in raw:
        if isinstance(entry, str):
            shots.append({"path": resolve_path(manifest_dir, entry), "caption": ""})
        elif isinstance(entry, dict) and "path" in entry:
            shots.append({
                "path": resolve_path(manifest_dir, entry["path"]),
                "caption": entry.get("caption", ""),
            })
        else:
            print(f"warning: skipping unrecognised screenshot entry: {entry!r}", file=sys.stderr)
    return shots


# ---------------------------------------------------------------------------
# Markdown → docx converter
# ---------------------------------------------------------------------------

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
CODE_SIZE = Pt(9)
BODY_SIZE = Pt(11)

# Shading colour for code blocks (light grey).
CODE_FILL = "F2F2F2"
# Border colour for code blocks.
CODE_BORDER = "D9D9D9"


def _set_cell_shading(cell, color: str) -> None:
    """Set the background shading of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = tcPr.find(qn("w:shd"))
    if shading is None:
        shading = tcPr.makeelement(qn("w:shd"), {})
        tcPr.append(shading)
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")


def _add_paragraph_with_text(doc_or_parent, text: str, style: str | None = None):
    """Add a paragraph, returning the paragraph object."""
    p = doc_or_parent.add_paragraph(style=style)
    if text:
        p.add_run(text)
    return p


class MarkdownToDocx:
    """Converts a markdown string into a python-docx Document."""

    def __init__(self, document: Document, screenshots: list[dict[str, str]]):
        self.doc = document
        self.screenshots = screenshots
        self._shot_index = 0  # for embedding in order

        # Ensure key styles exist and tweak them.
        self._ensure_styles()

    # -- style setup ----------------------------------------------------------

    def _ensure_styles(self) -> None:
        """Tweak built-in styles so headings and body look clean."""
        # Normal body
        style = self.doc.styles["Normal"]
        style.font.name = BODY_FONT
        style.font.size = BODY_SIZE
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(0)

        for level, size in [(1, 16), (2, 14), (3, 12)]:
            hstyle = self.doc.styles[f"Heading {level}"]
            hstyle.font.name = BODY_FONT
            hstyle.font.size = Pt(size)
            hstyle.font.bold = True
            hstyle.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            hstyle.paragraph_format.space_before = Pt(12)
            hstyle.paragraph_format.space_after = Pt(4)

    # -- top-level parser -----------------------------------------------------

    def convert(self, markdown: str) -> None:
        """Parse markdown line-by-line and build the document."""
        lines = markdown.splitlines()

        # We accumulate block-level tokens.
        i = 0
        while i < len(lines):
            line = lines[i]

            # Empty line — skip.
            if line.strip() == "":
                i += 1
                continue

            # Code block (fenced).
            if line.strip().startswith("```"):
                i, code = self._read_code_block(lines, i)
                self._emit_code_block(code)
                continue

            # Heading
            heading_match = re.match(r"^(#{1,6})\s+(.*)", line)
            if heading_match:
                level = min(len(heading_match.group(1)), 3)  # cap at H3
                text = heading_match.group(2).strip()
                self.doc.add_heading(text, level=level)
                i += 1
                continue

            # Unordered list item.
            list_match = re.match(r"^(\s*)[-*+]\s+(.*)", line)
            if list_match:
                i, items = self._read_list(lines, i)
                for item_text in items:
                    p = self.doc.add_paragraph(style="List Bullet")
                    self._add_formatted_runs(p, item_text)
                continue

            # Ordered list item.
            ordered_match = re.match(r"^(\s*)\d+\.\s+(.*)", line)
            if ordered_match:
                i, items = self._read_ordered_list(lines, i)
                for item_text in items:
                    p = self.doc.add_paragraph(style="List Number")
                    self._add_formatted_runs(p, item_text)
                continue

            # Horizontal rule
            if re.match(r"^[-*_]{3,}\s*$", line):
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(8)
                pPr = p._p.get_or_add_pPr()
                pBdr = pPr.makeelement(qn("w:pBdr"), {})
                bottom = pBdr.makeelement(qn("w:bottom"), {
                    qn("w:val"): "single",
                    qn("w:sz"): "6",
                    qn("w:space"): "1",
                    qn("w:color"): "CCCCCC",
                })
                pBdr.append(bottom)
                pPr.append(pBdr)
                i += 1
                continue

            # Inline image directive: ![caption](path) on its own line.
            img_match = re.match(r"^!\[(.*)\]\((.+)\)$", line.strip())
            if img_match:
                caption = img_match.group(1) or ""
                alt_path = img_match.group(2)
                # Try to resolve against known screenshots by basename.
                self._embed_image_by_name(alt_path, caption)
                i += 1
                continue

            # Regular paragraph.
            p = self.doc.add_paragraph()
            self._add_formatted_runs(p, line)
            i += 1

        # -- Insert table of contents at the very start -------------------
        self._insert_toc()

    # -- block helpers --------------------------------------------------------

    def _read_code_block(self, lines: list[str], start: int) -> tuple[int, str]:
        """Read a fenced code block. Returns (next_line_index, code_text)."""
        # Skip the opening fence; extract optional language.
        i = start + 1
        code_lines: list[str] = []
        while i < len(lines):
            if lines[i].strip().startswith("```"):
                i += 1
                break
            code_lines.append(lines[i])
            i += 1
        return i, "\n".join(code_lines)

    def _read_list(self, lines: list[str], start: int) -> tuple[int, list[str]]:
        """Read consecutive bullet list items."""
        items: list[str] = []
        i = start
        while i < len(lines):
            m = re.match(r"^(\s*)[-*+]\s+(.*)", lines[i])
            if not m:
                break
            items.append(m.group(2))
            i += 1
        return i, items

    def _read_ordered_list(self, lines: list[str], start: int) -> tuple[int, list[str]]:
        """Read consecutive ordered list items."""
        items: list[str] = []
        i = start
        while i < len(lines):
            m = re.match(r"^(\s*)\d+\.\s+(.*)", lines[i])
            if not m:
                break
            items.append(m.group(2))
            i += 1
        return i, items

    # -- inline formatting ----------------------------------------------------

    def _add_formatted_runs(self, paragraph, text: str) -> None:
        """Parse inline markdown (bold, italic, code) and add runs."""
        # Pattern matches: **bold**, *italic*, `code`
        pattern = re.compile(
            r"(\*\*(.+?)\*\*)|"     # bold
            r"(\*(.+?)\*)|"          # italic
            r"(`(.+?)`)"             # inline code
        )

        last = 0
        for m in pattern.finditer(text):
            # Plain text before this match.
            if m.start() > last:
                paragraph.add_run(text[last:m.start()])

            if m.group(1):  # bold
                run = paragraph.add_run(m.group(2))
                run.bold = True
            elif m.group(3):  # italic
                run = paragraph.add_run(m.group(4))
                run.italic = True
            elif m.group(5):  # inline code
                run = paragraph.add_run(m.group(6))
                run.font.name = CODE_FONT
                run.font.size = CODE_SIZE

            last = m.end()

        # Trailing plain text.
        if last < len(text):
            paragraph.add_run(text[last:])

    # -- code blocks ----------------------------------------------------------

    def _emit_code_block(self, code: str) -> None:
        """Insert a code block as a single-cell table with shading."""
        if not code.strip():
            return

        table = self.doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.cell(0, 0)

        # Remove default empty paragraph.
        for p in cell.paragraphs:
            p.clear()

        # Add each line of code.
        for line_no, line_text in enumerate(code.splitlines()):
            if line_no == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15

            run = p.add_run(line_text if line_text else " ")
            run.font.name = CODE_FONT
            run.font.size = CODE_SIZE

        _set_cell_shading(cell, CODE_FILL)

        self.doc.add_paragraph()  # spacer

    # -- images ---------------------------------------------------------------

    def _embed_image_by_name(self, name_or_path: str, caption: str) -> None:
        """Embed a screenshot, matching by basename or path."""
        # If it's an absolute path that exists, use it directly.
        if os.path.isfile(name_or_path):
            self._insert_image(name_or_path, caption)
            return

        # Try to find among manifest screenshots.
        base = os.path.basename(name_or_path)
        for shot in self.screenshots:
            if os.path.basename(shot["path"]) == base or shot["path"] == name_or_path:
                self._insert_image(shot["path"], caption or shot.get("caption", ""))
                return

        # If there are unplaced screenshots, use the next one in order.
        if self._shot_index < len(self.screenshots):
            shot = self.screenshots[self._shot_index]
            self._insert_image(shot["path"], caption or shot.get("caption", ""))
            self._shot_index += 1
            return

        # Fallback: try the path as-is.
        print(f"warning: screenshot not found: {name_or_path}", file=sys.stderr)

    def _insert_image(self, image_path: str, caption: str) -> None:
        """Insert an inline image into the document."""
        if not os.path.isfile(image_path):
            print(f"warning: image file not found: {image_path}", file=sys.stderr)
            return

        try:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=Inches(5.5))
        except Exception as exc:
            print(f"warning: could not embed {image_path}: {exc}", file=sys.stderr)
            return

        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.style = self.doc.styles["Normal"]
            cap_run = cap.add_run(caption)
            cap_run.italic = True
            cap_run.font.size = Pt(9)
            cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        self.doc.add_paragraph()  # spacer

    # -- table of contents ----------------------------------------------------

    def _insert_toc(self) -> None:
        """Insert a Word table-of-contents field at the start of the document."""
        # Move to the beginning: insert after the first paragraph (which is empty
        # in a fresh document) or create one.
        if len(self.doc.paragraphs) == 0:
            self.doc.add_paragraph()

        first = self.doc.paragraphs[0]

        # Insert a TOC paragraph before the first content paragraph.
        toc_para = first.insert_paragraph_before("Table of Contents")
        toc_para.style = self.doc.styles["Heading 1"]

        # Insert the actual TOC field.
        toc_field_para = toc_para.insert_paragraph_before()

        run = toc_field_para.add_run()
        fldChar_begin = run._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "begin"})
        run._r.append(fldChar_begin)

        run2 = toc_field_para.add_run()
        instrText = run2._r.makeelement(qn("w:instrText"), {})
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        run2._r.append(instrText)

        run3 = toc_field_para.add_run()
        fldChar_separate = run3._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "separate"})
        run3._r.append(fldChar_separate)

        run4 = toc_field_para.add_run("(Right-click and select 'Update Field' to populate the table of contents.)")
        run4.font.italic = True
        run4.font.size = Pt(9)
        run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        run5 = toc_field_para.add_run()
        fldChar_end = run5._r.makeelement(qn("w:fldChar"), {qn("w:fldCharType"): "end"})
        run5._r.append(fldChar_end)

        # Spacer after TOC.
        toc_field_para.insert_paragraph_before()


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Convert a docgen manifest into a .docx document."
    )
    parser.add_argument(
        "--manifest", "-m",
        required=True,
        help="Path to the manifest JSON file produced by the docgen tool.",
    )
    parser.add_argument(
        "--output", "-o",
        required=True,
        help="Output .docx file path.",
    )
    args = parser.parse_args()

    manifest_path = os.path.abspath(args.manifest)
    manifest_dir = os.path.dirname(manifest_path)

    data = load_manifest(manifest_path)

    md_path = resolve_path(manifest_dir, data["markdown"])
    if not os.path.isfile(md_path):
        sys.exit(f"markdown file not found: {md_path}")

    with open(md_path, "r", encoding="utf-8") as fh:
        markdown = fh.read()

    screenshots = normalise_screenshots(data.get("screenshots"), manifest_dir)

    doc = Document()

    # Page margins.
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    converter = MarkdownToDocx(doc, screenshots)
    converter.convert(markdown)

    output_path = os.path.abspath(args.output)
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
